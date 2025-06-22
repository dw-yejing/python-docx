from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, nsmap

# 创建新文档
document = Document()

# 添加标题
document.add_heading('Document Title', 0)

# 添加段落
p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

# 添加表格
table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'

table.style = 'TableGrid' 

# 保存文档
document.save('demo.docx')

